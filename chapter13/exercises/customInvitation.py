import docx, re

guestFile = open('guests.txt')
doc = docx.Document('.//base.docx')
regex = re.compile(r'(.*)\n')
content = guestFile.readlines()
for name in content:
    isLast = content.index(name) == len(content) - 1
    res = regex.match(name)
    if res:
        name = res[1]
    style = 'custom'
    styleName = 'name'
    doc.add_paragraph('It would be a pleasure to have the company of', style)
    pName = doc.add_paragraph(name, style)
    pName.style = styleName
    p = doc.add_paragraph('', style)
    r = p.add_run('At')
    r.underline = True
    p.add_run(' 11010 Memory Lane on the Evening of')
    doc.add_paragraph('April 1st', styleName)
    p2 = doc.add_paragraph('', style)
    r2 = p2.add_run('At')
    r2.underline = True
    p2.add_run(' 7 o\'clock')
    if not isLast:
        doc.add_page_break()

doc.save('test.docx')

guestFile.close()
